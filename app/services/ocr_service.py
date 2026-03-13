import csv
import os
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import List

from sqlalchemy.orm import Session

from app.core.config import POPPLER_PATH, TESSERACT_CMD
from app.core.errors import DependencyMissingError
from app.models.document_model import Document


@dataclass
class PageOcrResult:
    page_number: int
    text: str


def _resolve_tesseract_cmd() -> str | None:
    configured = TESSERACT_CMD.strip() if TESSERACT_CMD else ""
    if configured and Path(configured).exists():
        return configured
    return shutil.which("tesseract")


def _resolve_poppler_path() -> str | None:
    configured = POPPLER_PATH.strip() if POPPLER_PATH else ""
    if configured and Path(configured).exists():
        return configured
    pdftoppm_path = shutil.which("pdftoppm")
    if not pdftoppm_path:
        return None
    return str(Path(pdftoppm_path).parent)


def _has_meaningful_text(text: str) -> bool:
    compact = "".join(text.split())
    return len(compact) >= 20


def _ocr_pdf_file_scanned(pdf_path: str) -> List[PageOcrResult]:
    try:
        from pdf2image import convert_from_path  # type: ignore
    except ModuleNotFoundError as exc:
        raise DependencyMissingError(
            "pdf2image is required for scanned PDF OCR",
            details=[{"dependency": "pdf2image"}],
        ) from exc

    try:
        import pytesseract  # type: ignore
    except ModuleNotFoundError as exc:
        raise DependencyMissingError(
            "pytesseract is required for scanned PDF OCR",
            details=[{"dependency": "pytesseract"}],
        ) from exc

    tesseract_cmd = _resolve_tesseract_cmd()
    if not tesseract_cmd:
        raise DependencyMissingError(
            "Tesseract binary not found for scanned PDF OCR",
            details=[{"dependency": "tesseract"}],
        )
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    poppler_path = _resolve_poppler_path()
    convert_kwargs = {"dpi": 300}
    if poppler_path:
        convert_kwargs["poppler_path"] = poppler_path

    images = convert_from_path(pdf_path, **convert_kwargs)
    language = os.getenv("TESSERACT_LANG", "eng+vie")

    results: List[PageOcrResult] = []
    for idx, image in enumerate(images, start=1):
        text = pytesseract.image_to_string(image, lang=language) or ""
        results.append(PageOcrResult(page_number=idx, text=text))

    return results


def ocr_pdf_file(pdf_path: str) -> List[PageOcrResult]:
    """
    Extract text from PDF pages using PyPDF and fall back to OCR for scanned PDFs.
    """
    try:
        from pypdf import PdfReader  # type: ignore
    except ModuleNotFoundError as exc:
        raise DependencyMissingError(
            "pypdf is required to extract text from PDFs",
            details=[{"dependency": "pypdf"}],
        ) from exc

    reader = PdfReader(pdf_path)
    page_results: List[PageOcrResult] = []

    for idx, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        page_results.append(PageOcrResult(page_number=idx, text=text))

    if any(_has_meaningful_text(page.text) for page in page_results):
        return page_results

    return _ocr_pdf_file_scanned(pdf_path)


def _read_text_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as handle:
        return handle.read()


def _read_csv_file(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore", newline="") as handle:
        reader = csv.reader(handle)
        lines = [", ".join(row) for row in reader]
    return "\n".join(lines)


def _read_docx_file(file_path: str) -> str:
    try:
        import docx  # type: ignore
    except ModuleNotFoundError as exc:
        raise DependencyMissingError(
            "python-docx is required to extract text from DOCX files",
            details=[{"dependency": "python-docx"}],
        ) from exc

    document = docx.Document(file_path)
    lines = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells if cell.text]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _read_xlsx_file(file_path: str) -> str:
    try:
        import openpyxl  # type: ignore
    except ModuleNotFoundError as exc:
        raise DependencyMissingError(
            "openpyxl is required to extract text from XLSX files",
            details=[{"dependency": "openpyxl"}],
        ) from exc

    workbook = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
    lines = []
    for sheet in workbook.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value) for value in row if value is not None]
            if values:
                lines.append("\t".join(values))
    return "\n".join(lines)


def extract_document_text(document: Document) -> str:
    content_type = (document.content_type or "").lower()
    ext = Path(document.file_path).suffix.lower()

    if content_type == "application/pdf" or ext == ".pdf":
        pages = ocr_pdf_file(document.file_path)
        return "\n\n".join(f"[Page {p.page_number}]\n{p.text}" for p in pages)
    if content_type == "text/plain" or ext == ".txt":
        return _read_text_file(document.file_path)
    if content_type in {"text/csv", "application/vnd.ms-excel"} or ext == ".csv":
        return _read_csv_file(document.file_path)
    if (
        content_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or ext == ".docx"
    ):
        return _read_docx_file(document.file_path)
    if (
        content_type
        == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        or ext == ".xlsx"
    ):
        return _read_xlsx_file(document.file_path)

    raise ValueError(f"Unsupported content type: {document.content_type}")


def process_document_ocr(db: Session, document: Document):
    started_at = datetime.now(timezone.utc)
    document.processing_started_at = started_at
    document.status = "processing"
    document.processing_step = "ocr"
    document.processing_progress = 0
    document.processing_completed_at = None
    document.processing_duration_ms = None
    document.last_error = None
    db.commit()

    try:
        full_text = extract_document_text(document)

        # Save result
        document.text_content = full_text
        document.status = "processing"
        document.processing_step = "ocr"
        document.processing_progress = 35
        finished_at = datetime.now(timezone.utc)
        document.processing_duration_ms = int(
            (finished_at - started_at).total_seconds() * 1000
        )

    except Exception as e:
        document.status = "error"
        document.error_count = (document.error_count or 0) + 1
        document.last_error = str(e)
        document.processing_progress = document.processing_progress or 0
        document.processing_step = "error"
        document.processing_completed_at = datetime.now(timezone.utc)
        document.processing_duration_ms = int(
            (document.processing_completed_at - started_at).total_seconds() * 1000
        )

    db.commit()
    db.refresh(document)
    return document
